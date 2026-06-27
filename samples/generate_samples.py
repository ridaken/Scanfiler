"""Regenerate the sample input documents in samples/inbox/.

These committed samples let anyone try scanfiler locally without supplying their own
documents. Run from the repo root:

    python samples/generate_samples.py

Then point scanfiler at them:

    scanfiler -c samples/config.yaml plan --proposals samples/proposals.jsonl
"""

from __future__ import annotations

from pathlib import Path

INBOX = Path(__file__).resolve().parent / "inbox"


def _receipt_pdf(path: Path) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    lines = [
        "RIVERSIDE AUTO SERVICE",
        "123 Main St, Springfield",
        "",
        "RECEIPT  #  4471",
        "Date: 2025-06-15",
        "",
        "Kia Telluride - Brake pad replacement (front)   $189.00",
        "Synthetic oil change                            $ 64.00",
        "Shop supplies                                   $ 12.50",
        "",
        "TOTAL                                           $265.50",
        "Paid: VISA ****1234",
        "Thank you for your business!",
    ]
    page.insert_text((72, 90), "\n".join(lines), fontsize=11)
    doc.save(str(path))
    doc.close()


def _invoice_docx(path: Path) -> None:
    import docx

    d = docx.Document()
    d.add_heading("INVOICE", level=1)
    d.add_paragraph("Bright Spark Electric LLC")
    d.add_paragraph("Invoice #2025-0098    Date: 2025-05-02")
    d.add_paragraph("Bill to: Jordan Avery")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Panel upgrade to 200A"
    t.rows[0].cells[1].text = "$1,450.00"
    d.add_paragraph("Amount due: $1,450.00  -  Net 30")
    d.save(str(path))


def _drawing_image(path: Path) -> None:
    from PIL import Image, ImageDraw

    im = Image.new("RGB", (600, 400), (255, 252, 240))
    draw = ImageDraw.Draw(im)
    # A child's crayon-style house + sun, no text -> exercises the VLM / low-confidence path.
    draw.rectangle([180, 200, 380, 340], outline=(60, 90, 200), width=6)
    draw.polygon([(165, 200), (280, 110), (395, 200)], outline=(200, 60, 60), width=6)
    draw.rectangle([250, 270, 310, 340], outline=(60, 150, 60), width=5)
    draw.ellipse([470, 40, 560, 130], outline=(240, 190, 40), width=6)
    im.save(str(path))


def main() -> None:
    INBOX.mkdir(parents=True, exist_ok=True)
    _receipt_pdf(INBOX / "SCAN00001.pdf")
    _invoice_docx(INBOX / "SCAN00002.docx")
    _drawing_image(INBOX / "PIC00001.png")
    print(f"Wrote samples to {INBOX}")


if __name__ == "__main__":
    main()
