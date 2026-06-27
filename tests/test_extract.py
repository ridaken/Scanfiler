from pathlib import Path

from scanfiler.config import ExtractionConfig
from scanfiler.extract import classify, extract


def test_classify():
    assert classify(Path("a.pdf")) == "pdf"
    assert classify(Path("a.DOCX")) == "docx"
    assert classify(Path("a.JPG")) == "image"
    assert classify(Path("a.xyz")) == "unknown"


def test_unknown_type_returns_error(tmp_path):
    f = tmp_path / "weird.xyz"
    f.write_text("hello", encoding="utf-8")
    r = extract(f, ExtractionConfig())
    assert r.kind == "unknown"
    assert r.error is not None
    assert r.has_content is False


def test_corrupt_pdf_surfaces_error(tmp_path):
    f = tmp_path / "broken.pdf"
    f.write_bytes(b"%PDF-1.4 not really a pdf")
    r = extract(f, ExtractionConfig())
    assert r.error is not None  # exception captured, not raised


def test_pdf_text_mode_carries_no_images(tmp_path):
    import fitz

    f = tmp_path / "doc.pdf"
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), "Plenty of words " * 30)
    doc.save(str(f))
    doc.close()

    r = extract(f, ExtractionConfig(send_mode="text"))
    assert r.text.strip()
    assert r.images == []


def test_pdf_vision_mode_renders_images(tmp_path):
    import fitz

    f = tmp_path / "doc.pdf"
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), "x")
    doc.save(str(f))
    doc.close()

    r = extract(f, ExtractionConfig(send_mode="vision", pdf_max_pages=1))
    assert len(r.images) == 1


def test_image_downscaled_and_pngified(tmp_path):
    from PIL import Image

    f = tmp_path / "big.png"
    Image.new("RGB", (4000, 100), (1, 2, 3)).save(str(f))
    r = extract(f, ExtractionConfig())
    assert r.kind == "image"
    assert len(r.images) == 1
    # re-open the produced PNG; longest edge must be capped
    import io

    with Image.open(io.BytesIO(r.images[0])) as im:
        assert max(im.size) <= 1600


def test_docx_extracts_paragraphs_and_tables(tmp_path):
    import docx

    f = tmp_path / "d.docx"
    d = docx.Document()
    d.add_paragraph("Hello world")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Key"
    t.rows[0].cells[1].text = "Value"
    d.save(str(f))

    r = extract(f, ExtractionConfig())
    assert "Hello world" in r.text
    assert "Key | Value" in r.text
