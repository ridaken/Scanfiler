"""docx text extraction via python-docx (paragraphs + table cells)."""

from __future__ import annotations

from pathlib import Path

from . import ExtractResult

# Cap the text we carry forward; the model only needs enough to identify the doc.
_MAX_CHARS = 8000


def extract_docx(path: Path) -> ExtractResult:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()[:_MAX_CHARS]
    return ExtractResult(kind="docx", text=text)
